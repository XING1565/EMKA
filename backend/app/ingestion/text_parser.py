from __future__ import annotations

from io import BytesIO
from pathlib import Path


def parse_text_file(filename: str, content: bytes, metadata: dict) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".txt", ".md"}:
        return _decode_text(content)
    if suffix == ".pdf":
        return _parse_pdf(filename, content, metadata)
    if suffix == ".docx":
        return _parse_docx(filename, content, metadata)
    return _decode_text(content)


def _decode_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _parse_pdf(filename: str, content: bytes, metadata: dict) -> str:
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(content))
        metadata["page_count"] = len(reader.pages)
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"[page {index}]\n{text}")
        if pages:
            return "\n\n".join(pages)
        return f"No extractable PDF text found in {filename}."
    except Exception as exc:
        metadata["parse_error"] = str(exc)
        return f"PDF text extraction failed for {filename}: {exc}"


def _parse_docx(filename: str, content: bytes, metadata: dict) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_lines: list[str] = []
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    table_lines.append(f"table {table_index} row {row_index}: " + " | ".join(values))
        metadata["paragraph_count"] = len(paragraphs)
        metadata["table_count"] = len(document.tables)
        parts = paragraphs + table_lines
        if parts:
            return "\n".join(parts)
        return f"No extractable DOCX text found in {filename}."
    except Exception as exc:
        metadata["parse_error"] = str(exc)
        return f"DOCX text extraction failed for {filename}: {exc}"
